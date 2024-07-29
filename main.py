from pdf.PdfReader import PdfReader
from ai.Translator import Translator
from pdf.PageContent import PageContent
from docx import Document
from huggingface_hub import login

def main():
    counter = 1

    pdf_path = "C:/Users/moham/Downloads/questioning-the-influence-of-political-trust-on-individuals-unconventional-political-participation-in-palestine.pdf"
    output_docx_path = "C:/Users/moham/OneDrive/Desktop/Translator/translated_document.docx"

    try:
        pdf_reader = PdfReader(pdf_path)
        translator = Translator()
        document = Document()
        
        pages = pdf_reader.parse()
        
        for page in pages:
            print(f"Started with page {counter}")
            document.add_heading(f'Page {page.page_number}', level=1)
            document.add_paragraph(f'Header: {page.header}')
            
            for paragraph in page.paragraphs:
                if paragraph:  # Ensure the paragraph is not empty
                    translated_paragraph = translator.translate(paragraph)
                    document.add_paragraph(translated_paragraph)
            
            document.add_paragraph(f'Footer: {page.footer}')
            document.add_page_break()
            print(f"Done with page {counter}")
            counter += 1

        document.save(output_docx_path)
    except Exception as e:
        print(f"Error in main function: {e}")

if __name__ == "__main__":
    main()
